import os
from PyPDF2 import PdfReader
from docx import Document

def extracted_text(file_path):
    end = os.path.splitext(file_path)[1].lower()
    text = ""

    if end == ".pdf":
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    
    elif end == ".docx":
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + " "
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""

    elif end == ".txt":
        try:
            with open(file_path, "r", encoding="UTF-8") as f:
                text = f.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ""

    else:
        print("Unsupported format!")
        return ""
    
    return text.strip()